from typing import List, Dict, Tuple, Optional
from io import BytesIO
from docx import Document
import re
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def is_question_line(text: str) -> bool:
    """Intelligently determine if a line is a question using multiple heuristics."""
    text_lower = text.lower().strip()
    
    # Ignore empty text
    if not text or len(text.strip()) < 3:
        return False
    
    # Question indicators - strong signals
    question_patterns = [
        r'^[QR]\d+[\.:]?\s*',  # Q1, Q1., Q1:, R2.1, etc.
        r'^question\s+\d+',     # Question 1
        r'\?$',                 # Ends with question mark
        r'^what\s+',            # What is...
        r'^how\s+',             # How would...
        r'^why\s+',             # Why do...
        r'^when\s+',            # When did...
        r'^where\s+',           # Where is...
        r'^who\s+',             # Who is...
        r'^which\s+',           # Which of...
        r'^do\s+you\s+',        # Do you...
        r'^does\s+',            # Does it...
        r'^have\s+you\s+',      # Have you...
        r'^are\s+you\s+',       # Are you...
        r'^is\s+there\s+',      # Is there...
        r'^would\s+you\s+',     # Would you...
        r'^could\s+you\s+',     # Could you...
        r'^please\s+tell',      # Please tell me...
        r'^please\s+specify',   # Please specify...
        r'^please\s+describe',  # Please describe...
        r'^please\s+state',     # Please state...
        r'^please\s+provide',   # Please provide...
        r'^please\s+enter',     # Please enter...
        r'^record\s+the',       # Record the...
        r'^tell\s+me\s+',       # Tell me...
        r'^describe\s+',        # Describe...
        r'^specify\s+',         # Specify...
        r'^enter\s+',           # Enter...
        r'^write\s+',           # Write...
        r'^state\s+',           # State...
    ]
    
    for pattern in question_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    
    # Check for question-like structure even without question mark
    if len(text) > 10 and len(text) < 200:
        # Has question-like words
        question_words = ['what', 'how', 'why', 'when', 'where', 'who', 'which', 'do', 'does', 'did', 
                         'have', 'has', 'are', 'is', 'was', 'were', 'would', 'could', 'should']
        words = re.findall(r'\b\w+\b', text_lower)
        if any(word in question_words for word in words[:5]):  # Check first 5 words
            # Exclude common non-question patterns
            exclude_patterns = [
                r'^(note|instruction|guidance|section|for enumerator|do not read|for office use)',
                r'^the next',
                r'^you may',
                r'^this is',
                r'^read each',
                r'^thank you',
            ]
            if not any(re.match(pattern, text_lower) for pattern in exclude_patterns):
                return True
    
    return False


def is_note_or_instruction(text: str) -> bool:
    """Determine if text is a note, instruction, or header that should be treated as note type."""
    text_lower = text.lower().strip()
    
    # Strong indicators of notes/instructions
    note_patterns = [
        r'^(note|instruction|guidance|section)\s+',
        r'^for enumerator',
        r'^do not read',
        r'^for office use',
        r'^training only',
        r'^copyright',
        r'^disclaimer',
        r'^namaste',
        r'^interviewer',
        r'^the next questions? are',
        r'^you may select',
        r'^please note',
        r'^read each',
        r'^thank you',
        r'^this is',
        r'^we are',
        r'^recruitment and',
    ]
    
    for pattern in note_patterns:
        if re.match(pattern, text_lower):
            return True
    
    # Very long text blocks (>300 chars) without question marks are likely notes
    if len(text) > 300 and '?' not in text:
        return True
    
    # Section headers
    if re.match(r'^section\s+\d+[:]?', text_lower):
        return True
    
    return False


def is_instruction_line(text: str) -> bool:
    """Check if a line is an instruction/note that should be excluded from choices."""
    text_lower = text.lower().strip()
    
    instruction_patterns = [
        r'^(the next|you may|please note|read each|this is|we are|for enumerator)',
        r'select (all|more)',
        r'where applicable',
        r'open.?end',
        r'single code',
        r'multiple code',
        r'optional',
        r'required',
        r'please',
    ]
    
    # If line is too long, likely instruction
    if len(text) > 120:
        return True
    
    # If contains instruction keywords
    if any(re.search(pattern, text_lower) for pattern in instruction_patterns):
        return True
    
    return False


def extract_choices_intelligently(lines: List[str], question_text: str = "") -> List[str]:
    """Extract choices/options with intelligent pattern matching and context awareness."""
    choices = []
    cleaned_lines = []
    
    # First, clean and prepare all lines, excluding instructions
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip instruction lines
        if is_instruction_line(line):
            continue
            
        cleaned_lines.append(line)
    
    # Pattern 1: Standard numbered/bulleted choices
    # Matches: 1) Option, 1. Option, • Option, - Option, a) Option, A) Option, etc.
    choice_pattern = re.compile(
        r'^[\s]*(?:[-•*\u2022\u25CF\u2023]\s*)?'  # Bullet points
        r'(?:\(?\d+[\.\))]\s*)?'                   # Numbered: 1), 1., (1)
        r'(?:[a-zA-Z][\.\))]\s*)?'                 # Lettered: a), A)
        r'(.+?)$',                                  # The actual choice text
        re.IGNORECASE
    )
    
    for line in cleaned_lines:
        # Skip if line looks like a question (not a choice)
        if is_question_line(line):
            continue
        
        # Skip very long lines (likely not choices, but descriptions/notes)
        if len(line) > 150:
            continue
        
        # Skip lines ending with question mark (likely a question)
        if line.rstrip().endswith('?'):
            continue
        
        # Extract choice text using pattern
        match = choice_pattern.match(line)
        if match:
            choice_text = match.group(1).strip()
            # Clean up common prefixes and suffixes
            choice_text = re.sub(r'^[•\-\*]\s+', '', choice_text)
            choice_text = re.sub(r'\s+[–—\-]+$', '', choice_text)
            choice_text = choice_text.strip()
            
            # Exclude common markers that aren't actual choices
            exclude_markers = ['open-end', 'open end', 'single code', 'multiple code', 
                             'optional', 'required', 'specify', 'other (specify)']
            if any(marker in choice_text.lower() for marker in exclude_markers):
                continue
            
            if choice_text and len(choice_text) > 1:
                # Avoid duplicates (case-insensitive)
                if not any(c.lower() == choice_text.lower() for c in choices):
                    choices.append(choice_text)
        else:
            # If line doesn't match standard pattern but looks like a choice
            # (short, no question mark, not a question pattern, not an instruction)
            if (2 <= len(line) <= 100 and 
                not is_question_line(line) and 
                not is_instruction_line(line)):
                # Check if it's a valid choice candidate
                # Exclude common non-choice patterns
                if not re.match(r'^(note|instruction|section|for|the next|open.?end)', line.lower()):
                    # Exclude markers
                    if not any(marker in line.lower() for marker in ['open-end', 'single code', 'multiple code']):
                        if not any(c.lower() == line.lower() for c in choices):
                            choices.append(line)
    
    # Pattern 2: Extract inline choices from question text
    # Example: "What is your gender? (Male, Female, Other)"
    if not choices and question_text:
        # Look for choices in parentheses or after colon
        inline_patterns = [
            r'\(([^)]+)\)',  # (Choice1, Choice2, Choice3)
            r':\s*([^?]+?)(?:\?|$)',  # : Choice1, Choice2, Choice3
        ]
        
        for pattern in inline_patterns:
            matches = re.findall(pattern, question_text)
            for match in matches:
                # Split by commas, slashes, "or", "and"
                opts = re.split(r',|\/|\bor\b|\band\b', match)
                for opt in opts:
                    opt = opt.strip()
                    # Remove trailing punctuation
                    opt = re.sub(r'[.,;:]+$', '', opt)
                    if opt and 2 <= len(opt) <= 100:
                        if not any(c.lower() == opt.lower() for c in choices):
                            choices.append(opt)
    
    # Pattern 3: Infer Yes/No for binary questions (be more conservative)
    if not choices and question_text:
        # Only infer Yes/No for very clear binary questions
        # Strong binary indicators that suggest Yes/No response
        strong_binary_patterns = [
            r'\b(do you|does (it|this|that)|have you|are you|were you|did you)\s+',
            r'\bis there\b',
            r'\bare there\b',
        ]
        
        # Exclude questions that clearly need other responses
        exclude_patterns = [
            r'\b(what|which|how many|how much|how often|how long|when|where|who)\b',
            r'\b(specify|describe|explain|provide|state|enter|write|tell|list|name|select|choose)\b',
            r'\b(purpose|reason|type|kind|category|preference|option)\b',
        ]
        
        has_binary_indicator = any(re.search(pattern, question_text, re.IGNORECASE) for pattern in strong_binary_patterns)
        has_exclude_indicator = any(re.search(pattern, question_text, re.IGNORECASE) for pattern in exclude_patterns)
        
        # Only infer Yes/No if we have binary indicator and no exclude indicators
        if has_binary_indicator and not has_exclude_indicator:
            choices = ["Yes", "No"]
    
    return choices


def classify_question_type(question_text: str, choices: List[str], context_lines: List[str] = None) -> str:
    """Intelligently classify the question type based on multiple signals."""
    text_lower = question_text.lower()
    context_lower = " ".join([l.lower() for l in (context_lines or [])])
    combined_text = (text_lower + " " + context_lower).lower()
    
    # If it's a note/instruction, return note
    if is_note_or_instruction(question_text):
        return "note"
    
    # Check for select_multiple indicators
    select_multiple_keywords = [
        'select all that apply',
        'select all',
        'select more than one',
        'more than one',
        'multiple options',
        'multiple choices',
        'all that apply',
        'all applicable',
        'you may select',
        'tick all',
        'check all',
    ]
    
    is_multiple = any(keyword in combined_text for keyword in select_multiple_keywords)
    
    # If choices exist, determine between select_one and select_multiple
    if choices and len(choices) > 0:
        if is_multiple:
            return "select_multiple"
        else:
            return "select_one"
    
    # Check for integer/number type
    integer_keywords = [
        r'\b(how many|number of|count of|quantity of|total number|how much|amount of)\b',
        r'\bage\s+(in\s+)?years?\b',
        r'\byears?\s+old\b',
        r'\b\d+\s+to\s+\d+\s+years\b',  # Age ranges suggest integer
        r'\bphone\s+number\b',
        r'\bmobile\s+number\b',
        r'\bpin\s+code\b',
    ]
    
    if any(re.search(pattern, text_lower) for pattern in integer_keywords):
        return "integer"
    
    # Check for text/open-ended type
    text_keywords = [
        'specify',
        'describe',
        'explain',
        'provide details',
        'state',
        'enter',
        'write',
        'tell me',
        'your answer',
        'response:',
        'open-end',
        'open end',
        'free text',
        'comments',
        'remarks',
        'other (specify)',
        'please explain',
        'please specify',
        'please describe',
        'please state',
        'please provide',
        'please enter',
        'please write',
        'full name',
        'address',
        'city',
        'name',
    ]
    
    if any(keyword in text_lower for keyword in text_keywords):
        return "text"
    
    # Default to text if no clear indicator
    return "text"


def generate_field_name(text: str) -> str:
    """Generate a clean field name from question text using intelligent extraction."""
    # Remove question prefixes
    text = re.sub(r'^[QR]\d+[\.:]?\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'^question\s+\d+[\.:]?\s*', '', text, flags=re.IGNORECASE)
    
    # Remove common prefixes
    text = re.sub(r'^(what|how|why|when|where|who|which|do|does|did|are|is|was|were|have|has|would|could|should|please|tell|record)\s+', 
                  '', text, flags=re.IGNORECASE)
    
    # Remove parenthetical notes and markers
    text = re.sub(r'\([^)]*\)', '', text)  # Remove (SINGLE CODE), (Open-End), etc.
    text = re.sub(r'\[[^\]]*\]', '', text)
    
    # Remove question marks and special characters
    text = re.sub(r'[?.,\'":;!]', '', text)
    text = re.sub(r'[–—]', '-', text)
    
    # Extract meaningful words
    words = re.findall(r'[a-zA-Z0-9]+', text.lower())
    
    # Remove common stop words but keep important ones
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from'}
    words = [w for w in words if w not in stop_words]
    
    # Smart field name generation based on key terms
    field_name_map = {
        ('gender',): 'gender',
        ('age',): 'age',
        ('hair', 'type'): 'hair_type',
        ('hair', 'texture'): 'hair_texture',
        ('education',): 'education',
        ('internet', 'improve'): 'internet_improve',
        ('internet', 'purpose'): 'internet_purpose',
        ('internet', 'often'): 'internet_frequency',
        ('internet', 'freq'): 'internet_frequency',
        ('content', 'type'): 'content_type',
        ('free', 'time'): 'free_time_activities',
        ('activities',): 'activities',
        ('satisfaction',): 'satisfaction',
        ('satisf',): 'satisfaction',
        ('device',): 'device',
        ('habit',): 'habits',
        ('name', 'full'): 'full_name',
        ('participant', 'name'): 'participant_name',
        ('address', 'residential'): 'residential_address',
        ('phone',): 'phone',
        ('mobile',): 'mobile',
        ('city',): 'city',
        ('nccs',): 'nccs',
    }
    
    # Check for matching key terms
    words_set = set(words)
    for key_terms, field_name in field_name_map.items():
        if all(term in words_set for term in key_terms):
            return field_name
    
    # Fallback: use first 3-4 meaningful words, max 40 chars
    if words:
        field_name = '_'.join(words[:4])
        return field_name[:40] if len(field_name) <= 40 else '_'.join(words[:3])[:40]
    
    return 'field_unknown'


def extract_survey_units(docx_bytes: bytes) -> List[Dict]:
    """
    Main extraction function with intelligent parsing.
    Processes document paragraph by paragraph with context awareness.
    """
    doc = Document(BytesIO(docx_bytes))
    results = []
    
    # Get all paragraphs with their indices
    paragraphs = [(i, para.text.strip()) for i, para in enumerate(doc.paragraphs) if para.text.strip()]
    
    i = 0
    while i < len(paragraphs):
        para_idx, text = paragraphs[i]
        
        # Handle notes/instructions that appear standalone
        if is_note_or_instruction(text):
            # Check if it's a section header - always include
            if text.lower().startswith('section'):
                results.append({
                    "field_name": generate_field_name(text),
                    "question_text": text,
                    "choices": [],
                    "choice_list_name": "",
                    "type": "note",
                    "appearance": "",
                    "relevance": "",
                    "required": ""
                })
                i += 1
                continue
            # For other instructions, check if next item is a question
            # If next is question, this instruction is likely context for that question
            # Otherwise, treat as standalone note
            if i + 1 < len(paragraphs):
                next_idx, next_text = paragraphs[i + 1]
                if not is_question_line(next_text):
                    # Next isn't a question, so this might be a standalone note
                    if len(text) < 200 and '?' not in text:
                        results.append({
                            "field_name": generate_field_name(text),
                            "question_text": text,
                            "choices": [],
                            "choice_list_name": "",
                            "type": "note",
                            "appearance": "",
                            "relevance": "",
                            "required": ""
                        })
            # Skip standalone instructions (they'll be handled as context if needed)
            i += 1
            continue
        
        # If this looks like a question, start processing
        if is_question_line(text):
            question_text = text
            question_original = text
            
            # Look ahead to collect choices and context
            choices = []
            context_lines = []
            j = i + 1
            
            # Collect potential choices and context (up to next question or note)
            while j < len(paragraphs):
                next_idx, next_text = paragraphs[j]
                
                # Stop if we hit another question
                if is_question_line(next_text):
                    break
                
                # Stop if we hit a section header or major note
                if (is_note_or_instruction(next_text) and 
                    (next_text.lower().startswith('section') or len(next_text) > 200)):
                    break
                
                # Skip instruction lines that appear between questions
                if is_instruction_line(next_text):
                    # These are context but not choices - skip them from choice extraction
                    j += 1
                    continue
                
                # Collect as potential choice or context
                if not is_note_or_instruction(next_text):
                    # Check if it looks like a choice
                    if (len(next_text) <= 150 and 
                        not next_text.endswith('?') and
                        not is_question_line(next_text) and
                        not is_instruction_line(next_text)):
                        context_lines.append(next_text)
                    elif len(next_text) <= 200 and not is_instruction_line(next_text):
                        context_lines.append(next_text)
                
                j += 1
            
            # Extract choices from collected lines
            choices = extract_choices_intelligently(context_lines, question_text)
            
            # Classify the question type
            q_type = classify_question_type(question_text, choices, context_lines)
            
            # Generate field name
            field_name = generate_field_name(question_text)
            
            # Clean question text (remove parenthetical markers like (SINGLE CODE), (Open-End))
            clean_question = re.sub(r'\s*\([^)]*(?:SINGLE\s+CODE|MULTIPLE\s+CODE|Open-End|open-end|open\s+end)[^)]*\)', '', question_text, flags=re.IGNORECASE)
            # Remove placeholder underscores (___ or _ _ _)
            clean_question = re.sub(r'\s*_+\s*$', '', clean_question)
            # Remove trailing colons if question already has question mark
            if clean_question.rstrip().endswith('?'):
                clean_question = re.sub(r':\s*$', '', clean_question)
            elif not clean_question.rstrip().endswith('?'):
                # Add question mark if missing (unless it's a statement)
                if re.search(r'\b(what|how|why|when|where|who|which|do|does|are|is|have|would|could)', clean_question.lower()[:20]):
                    clean_question = re.sub(r':\s*$', '?', clean_question)
            # Clean up double question marks or punctuation
            clean_question = re.sub(r'\?+', '?', clean_question)
            clean_question = clean_question.strip()
            
            # If question has (Open-End) marker, force type to text and clear choices
            if re.search(r'\(open.?end\)', question_text, re.IGNORECASE):
                q_type = "text"
                choices = []
            
            # Determine required field
            required = "yes"
            if q_type == "select_multiple":
                required = "no"
            elif q_type == "note":
                required = ""
            elif "(optional)" in question_text.lower() or "(not required)" in question_text.lower():
                required = "no"
            
            results.append({
                "field_name": field_name,
                "question_text": clean_question if clean_question else question_text,
                "choices": choices,
                "choice_list_name": f"{field_name}_list" if choices else "",
                "type": q_type,
                "appearance": "",
                "relevance": "",
                "required": required
            })
            
            # Move to next unprocessed paragraph
            i = j
        else:
            # Not a clear question - might be a continuation, note, or we missed it
            # Try to extract as note if it seems important
            if len(text) < 300 and not any(r['question_text'] == text for r in results):
                # Check if it could be a missed question
                if '?' in text or any(word in text.lower()[:50] for word in ['what', 'how', 'which', 'do you', 'are you']):
                    # Retry as question
                    choices = extract_choices_intelligently([], text)
                    q_type = classify_question_type(text, choices)
                    
                    results.append({
                        "field_name": generate_field_name(text),
                        "question_text": text,
                        "choices": choices,
                        "choice_list_name": f"{generate_field_name(text)}_list" if choices else "",
                        "type": q_type,
                        "appearance": "",
                        "relevance": "",
                        "required": "yes" if q_type not in ["note", "select_multiple"] else ("no" if q_type == "select_multiple" else "")
                    })
            
            i += 1
    
    # Process tables
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not row_texts:
                continue
            
            # Heuristic: first cell is question, rest are choices
            question = row_texts[0]
            potential_choices = row_texts[1:] if len(row_texts) > 1 else []
            
            # Only process if first cell looks like a question
            if is_question_line(question) or '?' in question:
                choices = extract_choices_intelligently(potential_choices, question)
                q_type = classify_question_type(question, choices, potential_choices)
                
                clean_question = re.sub(r'\s*\([^)]*(?:SINGLE\s+CODE|MULTIPLE\s+CODE|Open-End)[^)]*\)', '', question)
                clean_question = clean_question.strip()
                
                results.append({
                    "field_name": generate_field_name(question),
                    "question_text": clean_question if clean_question else question,
                    "choices": choices,
                    "choice_list_name": f"{generate_field_name(question)}_list" if choices else "",
                    "type": q_type,
                    "appearance": "",
                    "relevance": "",
                    "required": "yes" if q_type not in ["note", "select_multiple"] else ("no" if q_type == "select_multiple" else "")
                })
    
    # Post-processing: clean up and merge orphaned choices
    # If a result has no choices but next result is just choices, merge them
    for i in range(len(results) - 1):
        if results[i]['type'] in ['select_one', 'select_multiple'] and not results[i]['choices']:
            # Check if next item could be choices
            next_item = results[i + 1]
            if (next_item['type'] == 'note' and 
                len(next_item['question_text']) < 100 and
                '?' not in next_item['question_text']):
                # Try to extract as choices
                potential_choices = extract_choices_intelligently([next_item['question_text']], results[i]['question_text'])
                if potential_choices:
                    results[i]['choices'] = potential_choices
                    results[i]['choice_list_name'] = f"{results[i]['field_name']}_list"
                    # Remove the orphaned note
                    results.pop(i + 1)
                    break
    
    logger.info(f"Extracted {len(results)} survey units")
    return results