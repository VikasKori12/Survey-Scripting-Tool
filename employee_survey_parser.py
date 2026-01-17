"""
Intelligent parser for employee survey documents.
Handles Likert-scale statements, demographic questions, and structured sections.
"""
from typing import List, Dict, Tuple, Optional
from io import BytesIO
from docx import Document
import re
import logging

logger = logging.getLogger(__name__)


def detect_scale_options(text: str, context: List[str] = None) -> Optional[List[str]]:
    """
    Detect scale options from text (e.g., "Strongly Disagree, Disagree, Neither, Agree, Strongly Agree")
    Returns list of scale options if found, None otherwise.
    """
    text_lower = text.lower()
    
    # Common scale patterns
    scale_patterns = [
        # 5-point agree/disagree scales
        (r'(strongly\s+disagree|strongly\s+disagree)',
         r'(somewhat\s+disagree|disagree|slightly\s+disagree)',
         r'(neither|neutral|neither\s+agree\s+nor\s+disagree)',
         r'(somewhat\s+agree|agree|slightly\s+agree)',
         r'(strongly\s+agree|strongly\s+agree)'),
        
        # 5-point frequency scales
        (r'(never)',
         r'(rarely|seldom)',
         r'(sometimes|occasionally)',
         r'(often|frequently)',
         r'(always|very\s+often)'),
        
        # 5-point quality scales
        (r'(very\s+poorly|very\s+poor|poorly)',
         r'(poor|below\s+average)',
         r'(average|about\s+average|moderate)',
         r'(well|good|above\s+average)',
         r'(very\s+well|excellent|very\s+good)'),
        
        # 5-point class scales
        (r'(lower\s+class)',
         r'(lower\s+middle\s+class)',
         r'(middle\s+class)',
         r'(upper\s+middle\s+class)',
         r'(upper\s+class)'),
    ]
    
    # Try to find scale in text
    for pattern_group in scale_patterns:
        matches = []
        for pattern in pattern_group:
            match = re.search(pattern, text_lower, re.IGNORECASE)
            if match:
                matches.append(match.group(0))
        if len(matches) >= 3:  # At least 3 matches found
            # Try to extract full scale from context
            return extract_scale_from_context(text, context)
    
    # Check if scale is mentioned (like "5-point scale")
    if re.search(r'\d+\s*[-]?\s*point\s+scale', text_lower):
        return extract_scale_from_context(text, context)
    
    return None


def extract_scale_from_context(text: str, context: List[str] = None) -> Optional[List[str]]:
    """Extract scale options from surrounding context."""
    # Common 5-point scales
    common_scales = {
        'agree': ['Strongly disagree', 'Somewhat disagree', 'Neither', 'Somewhat agree', 'Strongly agree'],
        'frequency': ['Never', 'Rarely', 'Sometimes', 'Often', 'Always'],
        'quality': ['Very poorly', 'Poorly', 'About average', 'Well', 'Very well'],
        'class': ['Lower class', 'Lower middle class', 'Middle class', 'Upper middle class', 'Upper class'],
    }
    
    text_lower = text.lower()
    
    # Check for scale type indicators
    if any(word in text_lower for word in ['agree', 'disagree', 'extent you agree']):
        return common_scales['agree']
    elif any(word in text_lower for word in ['often', 'frequently', 'never', 'always', 'rarely']):
        return common_scales['frequency']
    elif any(word in text_lower for word in ['groomed', 'well', 'poorly', 'quality']):
        return common_scales['quality']
    elif any(word in text_lower for word in ['class', 'socioeconomic', 'social class']):
        return common_scales['class']
    
    # Default to agree scale if scale is mentioned
    if 'scale' in text_lower or 'point' in text_lower:
        return common_scales['agree']
    
    return None


def is_scale_definition(text: str) -> bool:
    """Check if text is defining a scale (instructions before statements)."""
    text_lower = text.lower()
    
    scale_indicators = [
        'select the response',
        'circle the response',
        'point scale',
        'best reflects',
        'level of agreement',
        'extent you agree',
        'indicate to what extent',
        'select the response on the',
        'best describes',
    ]
    
    return any(indicator in text_lower for indicator in scale_indicators)


def is_statement(text: str) -> bool:
    """Check if text is a statement/question that uses a scale."""
    text_lower = text.lower().strip()
    
    # Skip if it's a scale definition or instruction
    if is_scale_definition(text_lower):
        return False
    
    # Skip if it's a section header
    if text_lower.startswith('section') or 'demographic' in text_lower:
        return False
    
    # Skip if it's an instruction line
    if any(word in text_lower for word in ['please select', 'please read', 'please indicate', 'please answer']):
        return False
    
    # Statements are usually:
    # - Complete sentences
    # - Not questions (usually)
    # - Not too short or too long
    if 10 <= len(text) <= 300:
        # Check if it looks like a statement (starts with "I", "My", "The", etc.)
        if re.match(r'^(I|My|The|This|Your|You|We|They|It)', text, re.IGNORECASE):
            return True
        # Or if it's a question about experiences/perceptions
        if any(word in text_lower for word in ['how', 'what', 'when', 'where', 'who']):
            if '?' in text:
                return True
    
    return False


def is_demographic_question(text: str) -> bool:
    """Check if text is a demographic question."""
    text_lower = text.lower()
    
    demographic_keywords = [
        'name', 'age', 'gender', 'education', 'occupation', 'marital status',
        'religion', 'income', 'employment', 'work experience', 'job title',
        'company', 'organization', 'household', 'family', 'parent', 'childhood',
        'birth', 'location', 'city', 'state', 'country', 'language', 'proficiency',
        'years', 'months', 'hours', 'level', 'position', 'status'
    ]
    
    # Check if it's asking for demographic info
    if any(keyword in text_lower for keyword in demographic_keywords):
        # But exclude statements
        if not is_statement(text):
            return True
    
    return False


def generate_field_name(text: str) -> str:
    """Generate a clean field name from text."""
    # Remove question prefixes
    text = re.sub(r'^[QR]\d+[\.:]?\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'^question\s+\d+[\.:]?\s*', '', text, flags=re.IGNORECASE)
    
    # Remove common prefixes
    text = re.sub(r'^(please|tell me|indicate|select|enter|write|provide|state|describe|specify)\s+', 
                  '', text, flags=re.IGNORECASE)
    
    # Remove parenthetical notes
    text = re.sub(r'\([^)]*\)', '', text)
    text = re.sub(r'\[[^\]]*\]', '', text)
    
    # Remove question marks and special characters
    text = re.sub(r'[?.,\'":;!]', '', text)
    text = re.sub(r'[–—]', '-', text)
    text = re.sub(r'_+', '', text)  # Remove underscores
    
    # Extract meaningful words
    words = re.findall(r'[a-zA-Z0-9]+', text.lower())
    
    # Remove stop words
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were'}
    words = [w for w in words if w not in stop_words and len(w) > 1]
    
    # Take first 4 meaningful words, max 40 chars
    if words:
        field_name = '_'.join(words[:4])
        return field_name[:40] if len(field_name) <= 40 else '_'.join(words[:3])[:40]
    
    return 'field_unknown'


def classify_demographic_type(text: str, choices: List[str] = None) -> str:
    """Classify demographic question type."""
    text_lower = text.lower()
    
    # Integer types
    if any(word in text_lower for word in ['age', 'years', 'months', 'hours', 'number', 'how many', 'how long']):
        return 'integer'
    
    # Text types
    if any(word in text_lower for word in ['name', 'specify', 'describe', 'provide', 'write', 'enter', 'address', 'occupation', 'title']):
        return 'text'
    
    # Select types (if choices exist)
    if choices:
        if any(word in text_lower for word in ['select all', 'all that apply', 'multiple']):
            return 'select_multiple'
        return 'select_one'
    
    # Default to text
    return 'text'


def extract_survey_units(docx_bytes: bytes) -> List[Dict]:
    """
    Main extraction function for employee surveys.
    Processes document to extract:
    1. Scale definitions (5-point scales)
    2. Statements using those scales
    3. Demographic questions
    """
    doc = Document(BytesIO(docx_bytes))
    results = []
    
    paragraphs = [(i, para.text.strip()) for i, para in enumerate(doc.paragraphs) if para.text.strip()]
    
    current_scale = None
    current_scale_name = None
    current_section = None
    i = 0
    
    while i < len(paragraphs):
        para_idx, text = paragraphs[i]
        
        # Detect section headers
        if text.lower().startswith('section') or 'demographic' in text.lower():
            current_section = text
            results.append({
                "field_name": generate_field_name(text),
                "question_text": text,
                "choices": [],
                "choice_list_name": "",
                "type": "note",
                "appearance": "",
                "relevance": "",
                "required": ""
            })
            i += 1
            continue
        
        # Detect scale definition
        if is_scale_definition(text):
            # Extract scale options
            context = [p[1] for p in paragraphs[max(0, i-2):min(len(paragraphs), i+5)]]
            scale_options = extract_scale_from_context(text, context)
            
            if scale_options:
                current_scale = scale_options
                # Generate scale name based on context
                if 'agree' in text.lower():
                    current_scale_name = 'agree'
                elif 'frequency' in text.lower() or 'often' in text.lower():
                    current_scale_name = 'frequency'
                elif 'groomed' in text.lower() or 'quality' in text.lower():
                    current_scale_name = 'rank'
                elif 'class' in text.lower():
                    current_scale_name = 'class'
                else:
                    current_scale_name = 'scale_' + str(len(results))
                
                # Add note for scale instruction
                results.append({
                    "field_name": generate_field_name(text),
                    "question_text": text,
                    "choices": [],
                    "choice_list_name": "",
                    "type": "note",
                    "appearance": "",
                    "relevance": "",
                    "required": ""
                })
            
            i += 1
            continue
        
        # Detect statements (that use current scale)
        if is_statement(text) and current_scale:
            field_name = generate_field_name(text)
            results.append({
                "field_name": field_name,
                "question_text": text,
                "choices": current_scale,
                "choice_list_name": current_scale_name,
                "type": "select_one",
                "appearance": "",
                "relevance": "",
                "required": "yes"
            })
            i += 1
            continue
        
        # Detect demographic questions
        if is_demographic_question(text):
            # Look ahead for choices
            choices = []
            j = i + 1
            while j < len(paragraphs) and j < i + 20:  # Look up to 20 lines ahead
                next_idx, next_text = paragraphs[j]
                
                # Stop if we hit another question or section
                if (is_demographic_question(next_text) or 
                    is_statement(next_text) or 
                    next_text.lower().startswith('section')):
                    break
                
                # Collect potential choices (short lines that aren't questions)
                if (10 <= len(next_text) <= 150 and 
                    not next_text.endswith('?') and
                    not is_scale_definition(next_text)):
                    # Check if it looks like a choice (not a statement)
                    if not re.match(r'^(I|My|The|This|Your|You|We)', next_text, re.IGNORECASE):
                        choices.append(next_text)
                
                j += 1
            
            # Classify type
            q_type = classify_demographic_type(text, choices)
            
            results.append({
                "field_name": generate_field_name(text),
                "question_text": text,
                "choices": choices if q_type in ['select_one', 'select_multiple'] else [],
                "choice_list_name": f"{generate_field_name(text)}_list" if choices and q_type in ['select_one', 'select_multiple'] else "",
                "type": q_type,
                "appearance": "",
                "relevance": "",
                "required": "yes" if q_type != "note" else ""
            })
            
            i = j  # Skip processed choices
            continue
        
        # Handle other text (notes, instructions)
        if len(text) > 5:
            # Check if it's an important note
            if any(word in text.lower() for word in ['note:', 'instruction', 'please note', 'important']):
                results.append({
                    "field_name": generate_field_name(text),
                    "question_text": text,
                    "choices": [],
                    "choice_list_name": "",
                    "type": "note",
                    "appearance": "",
                    "relevance": "",
                    "required": ""
                })
        
        i += 1
    
    # Process tables - statements are often in tables with scale options
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            continue
        
        # Check first row to understand table structure
        first_row = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]
        
        # If first row has scale options (like "Strongly disagree", "Somewhat disagree", etc.)
        # then this is a statement table
        first_row_lower = " ".join(first_row).lower()
        has_scale_in_first_row = any(word in first_row_lower for word in [
            'strongly', 'disagree', 'agree', 'somewhat', 'neither', 'never', 'always',
            'poorly', 'well', 'groomed', 'class'
        ])
        
        # Extract scale from table rows
        table_scale = None
        table_scale_name = None
        
        # Check first row for scale options
        if len(first_row) >= 5:
            # Check if first column is empty or "Statement" and rest are scale options
            first_col = first_row[0].lower() if first_row else ""
            is_header_row = (not first_col or first_col in ['statement', 'statements', 'item', 'items', ''] or len(first_col) < 10)
            
            if is_header_row:
                # Extract scale from remaining columns
                scale_options = []
                for cell_text in first_row[1:]:
                    cell_text = cell_text.strip().replace('\n', ' ').replace('\r', ' ')
                    if cell_text and len(cell_text) < 50:
                        scale_options.append(cell_text)
                
                if len(scale_options) >= 3:  # Valid scale
                    table_scale = scale_options
                    # Determine scale name
                    scale_text = " ".join(scale_options).lower()
                    if any('agree' in opt.lower() or 'disagree' in opt.lower() for opt in scale_options):
                        table_scale_name = 'agree'
                    elif any('groomed' in opt.lower() or 'well' in opt.lower() or 'poorly' in opt.lower() for opt in scale_options):
                        table_scale_name = 'rank'
                    elif any('class' in opt.lower() for opt in scale_options):
                        table_scale_name = 'class'
                    elif any('never' in opt.lower() or 'always' in opt.lower() or 'often' in opt.lower() for opt in scale_options):
                        table_scale_name = 'frequency'
                    else:
                        table_scale_name = f'scale_{table_idx}'
        
        # If scale not found in first row, check if first data row has scale in columns 2+
        if not table_scale and len(table.rows) > 1:
            second_row = [cell.text.strip() for cell in table.rows[1].cells if cell.text.strip()]
            if len(second_row) >= 5:
                # Check if columns 2+ look like scale options
                scale_candidates = []
                for cell_text in second_row[1:]:
                    cell_text = cell_text.strip().replace('\n', ' ').replace('\r', ' ')
                    if cell_text and len(cell_text) < 50:
                        scale_candidates.append(cell_text)
                
                # Verify these look like scale options
                if len(scale_candidates) >= 3:
                    scale_text_lower = " ".join(scale_candidates).lower()
                    if any(word in scale_text_lower for word in ['strongly', 'disagree', 'agree', 'somewhat', 'neither', 'never', 'always', 'well', 'poorly']):
                        table_scale = scale_candidates
                        if any('agree' in opt.lower() or 'disagree' in opt.lower() for opt in scale_candidates):
                            table_scale_name = 'agree'
                        else:
                            table_scale_name = f'scale_{table_idx}'
        
        # Process table rows
        start_row = 0
        
        # Process each row as a statement
        for row_idx in range(start_row, len(table.rows)):
            row = table.rows[row_idx]
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            
            if not row_texts:
                continue
            
            statement = row_texts[0]
            
            # Skip if statement is empty or looks like a header
            if not statement or len(statement) < 5:
                continue
            
            # Skip header rows
            if statement.lower() in ['statement', 'statements', 'item', 'items']:
                continue
            
            # Clean statement text
            statement = statement.replace('\n', ' ').replace('\r', ' ').strip()
            
            # Check if this row contains scale options (header row) - skip if we already have scale
            if not table_scale:
                row_texts_lower = " ".join(row_texts).lower()
                is_scale_header = any(word in row_texts_lower for word in [
                    'strongly', 'disagree', 'agree', 'somewhat', 'neither', 'never', 'always',
                    'poorly', 'well', 'groomed'
                ]) and len(row_texts) >= 5
                
                if is_scale_header:
                    # This row defines the scale
                    scale_options = []
                    for cell_text in row_texts[1:]:  # Skip first column
                        cell_text = cell_text.strip().replace('\n', ' ').replace('\r', ' ')
                        if cell_text and len(cell_text) < 50:
                            scale_options.append(cell_text)
                    if len(scale_options) >= 3:
                        table_scale = scale_options
                        if any('agree' in opt.lower() or 'disagree' in opt.lower() for opt in scale_options):
                            table_scale_name = 'agree'
                        elif any('groomed' in opt.lower() or 'well' in opt.lower() or 'poorly' in opt.lower() for opt in scale_options):
                            table_scale_name = 'rank'
                        else:
                            table_scale_name = f'scale_{table_idx}'
                    continue  # Skip this header row
            
            # If this looks like a statement and we have a scale
            if is_statement(statement) and table_scale:
                # Don't apply scale to demographic questions
                if not is_demographic_question(statement):
                    # Clean scale options
                    clean_scale = [opt.replace('\n', ' ').replace('\r', ' ').strip() for opt in table_scale]
                    results.append({
                        "field_name": generate_field_name(statement),
                        "question_text": statement,
                        "choices": clean_scale,
                        "choice_list_name": table_scale_name,
                        "type": "select_one",
                        "appearance": "",
                        "relevance": "",
                        "required": "yes"
                    })
                    continue
            
            if is_demographic_question(statement):
                # Demographic question in table
                potential_choices = row_texts[1:] if len(row_texts) > 1 else []
                q_type = classify_demographic_type(statement, potential_choices)
                results.append({
                    "field_name": generate_field_name(statement),
                    "question_text": statement,
                    "choices": potential_choices if q_type in ['select_one', 'select_multiple'] else [],
                    "choice_list_name": f"{generate_field_name(statement)}_list" if potential_choices and q_type in ['select_one', 'select_multiple'] else "",
                    "type": q_type,
                    "appearance": "",
                    "relevance": "",
                    "required": "yes" if q_type != "note" else ""
                })
            elif '?' in statement or len(statement) > 10:
                # Might be a question - try to extract choices from row
                potential_choices = row_texts[1:] if len(row_texts) > 1 else []
                if potential_choices and all(len(c) < 50 for c in potential_choices):
                    # Looks like choices
                    q_type = classify_demographic_type(statement, potential_choices)
                    results.append({
                        "field_name": generate_field_name(statement),
                        "question_text": statement,
                        "choices": potential_choices,
                        "choice_list_name": f"{generate_field_name(statement)}_list",
                        "type": "select_one",
                        "appearance": "",
                        "relevance": "",
                        "required": "yes"
                    })
                else:
                    # Text question
                    results.append({
                        "field_name": generate_field_name(statement),
                        "question_text": statement,
                        "choices": [],
                        "choice_list_name": "",
                        "type": "text",
                        "appearance": "",
                        "relevance": "",
                        "required": "yes"
                    })
    
    logger.info(f"Extracted {len(results)} survey units")
    return results
